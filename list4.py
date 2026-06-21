# mainKAA/views.py
from django.shortcuts import render, redirect
from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.decorators import login_required
from django.http import FileResponse
from django.core.files.base import ContentFile
from docx import Document
from io import BytesIO
import datetime
from .models import DocumentTemplate, GeneratedDocument, StudentProfile

def login_view(request):
    if request.method == 'POST':
        username = request.POST.get('username')
        password = request.POST.get('password')
        user = authenticate(request, username=username, password=password)
        if user is not None:
            login(request, user)
            StudentProfile.objects.get_or_create(user=user)
            return redirect('select_document')
        else:
            return render(request, 'login.html', {'error': 'Неверное имя пользователя или пароль'})
    return render(request, 'login.html')

def logout_view(request):
    logout(request)
    return redirect('login')

@login_required
def select_document(request):
    return render(request, 'select_document.html', {'user': request.user})

@login_required
def profile_view(request):
    profile, created = StudentProfile.objects.get_or_create(user=request.user)
    
    if request.method == 'POST':
        profile.fio = request.POST.get('fio', '')
        profile.grupa = request.POST.get('grupa', '')
        profile.kurs = request.POST.get('kurs') or None
        profile.obuch = request.POST.get('obuch', '')
        profile.spec = request.POST.get('spec', '')
        profile.vid = request.POST.get('vid', '')
        profile.kod = request.POST.get('kod', '')
        profile.adress = request.POST.get('adress', '')
        profile.ruka = request.POST.get('ruka', '')
        profile.save()
        return render(request, 'profile_form.html', {'profile': profile, 'saved': True})
    
    return render(request, 'profile_form.html', {'profile': profile})

def generate_docx(template, data):
    doc = Document(template.template_file)
    
    for paragraph in doc.paragraphs:
        for key, value in data.items():
            placeholder = f'{{{{{key}}}}}'
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, str(value))
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in data.items():
                    placeholder = f'{{{{{key}}}}}'
                    if placeholder in cell.text:
                        cell.text = cell.text.replace(placeholder, str(value))
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

@login_required
def document_form(request):
    template = DocumentTemplate.objects.first()
    
    if not template:
        return render(request, 'error.html', {'message': 'Нет загруженного шаблона документа'})
    
    profile, created = StudentProfile.objects.get_or_create(user=request.user)
    
    if request.method == 'POST':
        data = {
            'fio': request.POST.get('fio') or profile.fio or '',
            'grupa': request.POST.get('grupa') or profile.grupa or '',
            'kurs': request.POST.get('kurs') or profile.kurs or '',
            'obuch': request.POST.get('obuch') or profile.obuch or '',
            'spec': request.POST.get('spec') or profile.spec or '',
            'vid': request.POST.get('vid') or profile.vid or '',
            'kod': request.POST.get('kod') or profile.kod or '',
            'adress': request.POST.get('adress') or profile.adress or '',
            'ruka': request.POST.get('ruka') or profile.ruka or '',
            'data': request.POST.get('data', ''),
            'data2': request.POST.get('data2', ''),
            'god': request.POST.get('god', ''),
        }
        
        docx_content = generate_docx(template, data)
        safe_name = data['fio'].replace(' ', '_') if data['fio'] else 'document'
        filename = f"Attestatsionniy_list_{safe_name}_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        
        GeneratedDocument.objects.create(
            user=request.user,
            template=template,
            document_file=ContentFile(docx_content, name=filename)
        )
        
        return FileResponse(BytesIO(docx_content), as_attachment=True, filename=filename)
    
    return render(request, 'document_form.html', {
        'template': template,
        'profile': profile,
        'profile_exists': bool(profile.fio or profile.grupa or profile.kurs),
    })